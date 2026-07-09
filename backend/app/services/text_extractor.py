import io
import logging
import mimetypes
import zipfile
from pathlib import Path
from typing import Optional

from app.services.storage import storage

logger = logging.getLogger(__name__)


def _get_file_extension(object_key: str) -> str:
    """Return the lower-case file extension for the S3 object key."""
    return Path(object_key).suffix.lower()


def _is_docx_bytes(file_bytes: bytes) -> bool:
    """Detect whether the byte stream represents a DOCX archive."""
    try:
        return zipfile.is_zipfile(io.BytesIO(file_bytes))
    except Exception:
        return False


def _detect_file_type(file_bytes: bytes, object_key: str) -> str:
    """Detect file type using extension heuristics and file signature data."""
    extension = _get_file_extension(object_key)
    logger.debug("Detecting file type for object_key=%s extension=%s", object_key, extension)

    if extension == ".pdf" or file_bytes.startswith(b"%PDF-"):
        return "pdf"

    if extension == ".docx" or _is_docx_bytes(file_bytes):
        return "docx"

    if extension in {".txt", ".log", ".md"}:
        return "txt"

    if extension in {".jpg", ".jpeg", ".png"}:
        return "image"

    if file_bytes.startswith(b"\xff\xd8\xff") or file_bytes.startswith(b"\x89PNG"):
        return "image"

    guessed_type, _ = mimetypes.guess_type(object_key)
    if guessed_type:
        if guessed_type == "application/pdf":
            return "pdf"
        if guessed_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return "docx"
        if guessed_type.startswith("text/"):
            return "txt"
        if guessed_type.startswith("image/"):
            return "image"

    logger.debug("Unable to determine explicit file type; falling back to plain text")
    return "txt"


def _extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract readable text from a PDF document using PyMuPDF."""
    try:
        import fitz
    except ImportError:
        logger.warning("PDF extraction requires PyMuPDF but the package is not installed.")
        return ""

    try:
        with fitz.open(stream=file_bytes, filetype="pdf") as document:
            pages = [page.get_text("text") for page in document]
            return "\n".join(pages).strip()
    except Exception as exc:
        logger.exception("Failed to extract text from PDF: %s", exc)
        return ""


def _extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract readable text from a DOCX file.

    Prefer python-docx when installed, but fall back to a lightweight
    ZIP/XML parser so DOCX processing works without extra dependencies.
    """
    try:
        from docx import Document
        document = Document(io.BytesIO(file_bytes))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
        tables = []
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        tables.append(cell.text)
        return "\n".join(paragraphs + tables).strip()
    except Exception:
        logger.debug("python-docx unavailable or failed; using fallback DOCX extraction.")

    try:
        import zipfile
        import xml.etree.ElementTree as ET
    except Exception as exc:
        logger.exception("Failed to import fallback DOCX extraction dependencies: %s", exc)
        return ""

    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as archive:
            if "word/document.xml" not in archive.namelist():
                logger.warning("DOCX archive missing word/document.xml")
                return ""

            xml_bytes = archive.read("word/document.xml")
            root = ET.fromstring(xml_bytes)
            namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            paragraphs = []
            for paragraph in root.findall('.//w:p', namespaces):
                text_parts = [node.text for node in paragraph.findall('.//w:t', namespaces) if node.text]
                if text_parts:
                    paragraphs.append("".join(text_parts))
            return "\n".join(paragraphs).strip()
    except Exception as exc:
        logger.exception("Fallback DOCX extraction failed: %s", exc)
        return ""


def _extract_text_from_txt(file_bytes: bytes) -> str:
    """Extract readable text from a plain text file, using UTF-8 with replacement fallback."""
    try:
        return file_bytes.decode("utf-8").strip()
    except UnicodeDecodeError:
        logger.warning("UTF-8 decode failed for TXT file; using replacement characters.")
        return file_bytes.decode("utf-8", errors="replace").strip()
    except Exception as exc:
        logger.exception("Failed to extract text from TXT: %s", exc)
        return ""


def _placeholder_ocr(file_bytes: bytes, object_key: str) -> str:
    """Placeholder OCR: logs that OCR is not implemented yet and returns an empty string."""
    logger.debug("Placeholder OCR activated for image file '%s'. OCR is not implemented yet.", object_key)
    return ""


def extract_text(bucket_name: str, object_key: str) -> str:
    """Download a file from S3 and extract readable text from supported file formats."""
    logger.info("Starting text extraction for S3 object s3://%s/%s", bucket_name, object_key)

    try:
        file_bytes = storage.get_file_bytes({"bucket": bucket_name, "object_key": object_key})
        if not file_bytes:
            logger.warning("No file bytes returned from S3 for s3://%s/%s", bucket_name, object_key)
            return ""

        file_type = _detect_file_type(file_bytes, object_key)
        logger.debug("File type detected as %s for object_key=%s", file_type, object_key)

        if file_type == "pdf":
            return _extract_text_from_pdf(file_bytes)
        if file_type == "docx":
            return _extract_text_from_docx(file_bytes)
        if file_type == "txt":
            return _extract_text_from_txt(file_bytes)
        if file_type == "image":
            return _placeholder_ocr(file_bytes, object_key)

        logger.debug("Falling back to plain text extraction for object_key=%s", object_key)
        return _extract_text_from_txt(file_bytes)

    except Exception as exc:
        logger.exception("Text extraction failed for s3://%s/%s", bucket_name, object_key)
        return ""
